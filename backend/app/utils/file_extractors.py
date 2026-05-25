"""
Text extraction utilities for different file formats.

This module provides functions to extract text content from various file types
commonly used in document processing systems:

1. PDF (Portable Document Format)
   - Binary format widely used for documents
   - PyMuPDF (fitz) provides efficient page-by-page extraction
   - Preserves text layout and structure

2. DOCX (Microsoft Word)
   - XML-based format used by modern Microsoft Office
   - python-docx provides paragraph and table extraction
   - Can extract from paragraphs, tables, headers, footers

3. TXT (Plain Text)
   - Simplest format, already extractable
   - Just need to read with proper encoding detection

4. CSV (Comma-Separated Values)
   - Tabular data format
   - pandas provides efficient parsing and conversion to text

Why Text Extraction for RAG?
===========================
RAG (Retrieval-Augmented Generation) systems work by:
1. Splitting documents into chunks
2. Creating vector embeddings of chunks
3. Storing embeddings in vector database
4. At query time: retrieve relevant chunks, pass to LLM

Without text extraction, you can't:
- Split documents into meaningful chunks
- Create embeddings
- Enable semantic search over document content

Text extraction is the critical first step in the RAG pipeline.
"""

import csv
import io
from typing import Tuple

import fitz  # PyMuPDF
import pandas as pd
from docx import Document as DocxDocument


def extract_pdf_text(file_content: bytes) -> Tuple[str, int]:
    """
    Extract text from PDF file.
    
    PyMuPDF (fitz) Overview:
    - Lightweight, fast PDF processing library
    - Extracts text page by page
    - Returns text with layout preservation
    - Can extract metadata, images, etc.
    
    Why PyMuPDF over alternatives?
    - pdfplumber: Better for table extraction but slower
    - pypdf: Pure Python, slower than PyMuPDF
    - PyMuPDF: Fastest, C-based, most efficient
    
    Process:
    1. Open PDF from bytes in memory
    2. Iterate through each page
    3. Extract text from each page
    4. Join all pages with separator
    5. Clean up and return
    
    Args:
        file_content: PDF file as bytes
        
    Returns:
        Tuple of (extracted_text, character_count)
        
    Raises:
        Exception: If PDF is corrupted or cannot be opened
    """
    try:
        # Open PDF from bytes in memory (stream)
        # fitz.open(stream=bytes) opens from memory without disk write
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        
        text_content = []
        
        # Extract text from each page
        for page_num in range(len(pdf_document)):
            # Get page object
            page = pdf_document[page_num]
            
            # Extract text from page
            # PyMuPDF extracts text in reading order by default
            page_text = page.get_text()
            
            # Add page number marker for clarity in chunking
            if page_text.strip():
                text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        # Join all pages with clear separation
        extracted_text = "\n\n".join(text_content)
        
        # Close document
        pdf_document.close()
        
        return extracted_text, len(extracted_text)
    
    except Exception as e:
        raise Exception(f"Failed to extract PDF: {str(e)}")


def extract_docx_text(file_content: bytes) -> Tuple[str, int]:
    """
    Extract text from DOCX (Microsoft Word) file.
    
    python-docx Overview:
    - Reads DOCX as ZIP archive containing XML
    - Extracts from paragraphs, tables, headers, footers
    - Preserves document structure
    
    DOCX File Structure:
    - DOCX is actually a ZIP file
    - Contains XML files for document content
    - word/document.xml = main document content
    - word/styles.xml = formatting styles
    - word/header1.xml, word/footer1.xml = headers/footers
    - word/media/ = embedded images
    
    Process:
    1. Open DOCX from bytes (treated as file-like object)
    2. Extract paragraphs (main content)
    3. Extract tables (tabular content)
    4. Join all content
    5. Return full text
    
    Args:
        file_content: DOCX file as bytes
        
    Returns:
        Tuple of (extracted_text, character_count)
        
    Raises:
        Exception: If DOCX is corrupted or cannot be opened
    """
    try:
        # Open DOCX from bytes using BytesIO (file-like object)
        docx_file = io.BytesIO(file_content)
        document = DocxDocument(docx_file)
        
        text_content = []
        
        # Extract from paragraphs (main content)
        for para in document.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract from tables (tabular content)
        for table in document.tables:
            # Add table header marker
            text_content.append("\n--- Table ---")
            
            # Extract from each row
            for row in table.rows:
                # Extract text from each cell in row
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                )
                if row_text.strip():
                    text_content.append(row_text)
            
            text_content.append("--- End Table ---\n")
        
        # Join all content
        extracted_text = "\n".join(text_content)
        
        return extracted_text, len(extracted_text)
    
    except Exception as e:
        raise Exception(f"Failed to extract DOCX: {str(e)}")


def extract_txt_text(file_content: bytes) -> Tuple[str, int]:
    """
    Extract text from plain text file.
    
    Plain Text Handling:
    - Simplest format, already text
    - Just need to decode bytes to string
    - Handle encoding issues gracefully
    
    Encoding Strategy:
    - Try UTF-8 first (most common)
    - Fall back to latin-1 (never fails, catch-all)
    - Handle BOM (Byte Order Mark) if present
    
    Args:
        file_content: TXT file as bytes
        
    Returns:
        Tuple of (extracted_text, character_count)
        
    Raises:
        Exception: If file cannot be decoded
    """
    try:
        # Try UTF-8 first (most common encoding)
        try:
            text_content = file_content.decode("utf-8")
        except UnicodeDecodeError:
            # Fall back to latin-1 (covers most Western European)
            text_content = file_content.decode("latin-1")
        
        # Strip leading/trailing whitespace
        text_content = text_content.strip()
        
        return text_content, len(text_content)
    
    except Exception as e:
        raise Exception(f"Failed to extract TXT: {str(e)}")


def extract_csv_text(file_content: bytes) -> Tuple[str, int]:
    """
    Extract text from CSV (Comma-Separated Values) file.
    
    CSV Processing Strategy:
    - Parse CSV using pandas for robust handling
    - Convert tabular data to readable text
    - Preserve column structure and relationships
    
    Why pandas for CSV?
    - Handles different delimiters (comma, semicolon, tab)
    - Manages quoting and escaping properly
    - Detects encoding
    - Handles missing values
    
    Conversion to Text:
    - Format as pipe-separated for clarity: "Col1 | Col2 | Col3"
    - Each row on new line
    - Include column headers
    - Indicates table structure
    
    Args:
        file_content: CSV file as bytes
        
    Returns:
        Tuple of (extracted_text, character_count)
        
    Raises:
        Exception: If CSV is malformed or cannot be parsed
    """
    try:
        # Decode bytes to string
        # CSV is text-based, so decode first
        try:
            csv_string = file_content.decode("utf-8")
        except UnicodeDecodeError:
            csv_string = file_content.decode("latin-1")
        
        # Use StringIO to create file-like object from string
        csv_file = io.StringIO(csv_string)
        
        # Read CSV with pandas
        # pandas automatically detects delimiter
        df = pd.read_csv(csv_file)
        
        # Convert DataFrame to readable text format
        text_lines = []
        
        # Add header row
        header_row = " | ".join(str(col) for col in df.columns)
        text_lines.append(header_row)
        text_lines.append("-" * len(header_row))  # Separator
        
        # Add data rows
        for _, row in df.iterrows():
            row_text = " | ".join(str(val) for val in row.values)
            text_lines.append(row_text)
        
        # Join all lines
        extracted_text = "\n".join(text_lines)
        
        return extracted_text, len(extracted_text)
    
    except Exception as e:
        raise Exception(f"Failed to extract CSV: {str(e)}")


def extract_text_from_file(
    file_content: bytes,
    file_type: str
) -> Tuple[str, int]:
    """
    Extract text from any supported file type.
    
    Router function that dispatches to appropriate extractor
    based on MIME type.
    
    Supported MIME Types:
    - application/pdf: PDF documents
    - application/vnd.openxmlformats-officedocument.wordprocessingml.document: DOCX
    - text/plain: Plain text files
    - text/csv: CSV files
    
    Args:
        file_content: File content as bytes
        file_type: MIME type of file (e.g., "application/pdf")
        
    Returns:
        Tuple of (extracted_text, character_count)
        
    Raises:
        ValueError: If file type not supported
        Exception: If extraction fails
    """
    # Map MIME types to extractor functions
    extractors = {
        "application/pdf": extract_pdf_text,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_docx_text,
        "text/plain": extract_txt_text,
        "text/csv": extract_csv_text,
    }
    
    # Get extractor for this file type
    extractor = extractors.get(file_type)
    
    if not extractor:
        raise ValueError(
            f"Unsupported file type: {file_type}. "
            f"Supported types: {', '.join(extractors.keys())}"
        )
    
    # Call appropriate extractor
    return extractor(file_content)
